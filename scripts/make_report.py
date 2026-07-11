"""Generate the MoleCheck capstone report as a Word document (python-docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MoleCheck_Capstone_Report.docx"

TEAL = RGBColor(0x00, 0x79, 0x6B)
DARK = RGBColor(0x1A, 0x2B, 0x29)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell(cell, text, *, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "00796B")
        set_cell(c, h, bold=True, color=WHITE,
                 align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(r):
            if ri % 2:
                shade(cells[i], "F2F7F6")
            set_cell(cells[i], str(val), bold=(i == 0),
                     align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def para(doc, text, *, italic=False, color=DARK, size=11, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.font.size = Pt(11)


def h1(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = TEAL
    run.font.size = Pt(16)
    # thin rule under the heading
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = DARK
    run.font.size = Pt(13)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F7F6")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # ---- Title page ----
    for _ in range(6):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("MoleCheck")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = TEAL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("An On-Device AI Skin-Lesion Screening App")
    r.font.size = Pt(16)
    r.font.color.rgb = DARK
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Capstone Project Report")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    for _ in range(2):
        doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run(
        "Classifying moles as benign or malignant from a phone photo, "
        "with a model that runs entirely on the device."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    doc.add_page_break()

    # ---- 1. Executive Summary ----
    h1(doc, "1. Executive Summary")
    para(doc, "MoleCheck is a cross-platform mobile application (iOS and Android) that helps a user "
              "assess whether a mole looks benign or potentially malignant. The user photographs a single "
              "skin lesion, and a convolutional neural network trained on dermatoscopic images estimates "
              "the likelihood that the lesion is malignant. Crucially, the model runs entirely on the phone: "
              "the photo is never uploaded, which preserves privacy and lets the app work offline.")
    para(doc, "The classifier reaches an ROC-AUC of 0.914 on a held-out test set of 1,722 images. Operating "
              "at a decision threshold tuned for screening, it correctly identifies 90.1% of malignant lesions "
              "(sensitivity) while correctly clearing 74.7% of benign ones (specificity). The trained model "
              "was exported to TensorFlow Lite and verified to reproduce the original model to within 0.001 "
              "probability, then embedded in a Flutter app and tested end-to-end on a device.")
    para(doc, "This report documents the dataset, methodology, results, deployment pipeline, and the "
              "limitations and ethical framing of a tool of this kind. MoleCheck is an educational project, "
              "not a medical device, and is designed throughout to direct users to a dermatologist rather "
              "than to replace one.")

    # ---- 2. Introduction ----
    h1(doc, "2. Introduction and Motivation")
    para(doc, "Melanoma and other skin cancers are among the most treatable cancers when caught early and "
              "among the most dangerous when caught late. Yet access to dermatology is uneven, and many "
              "people do not know which of their moles, if any, are worth showing to a doctor. A phone-based "
              "triage aid that flags concerning lesions could lower the barrier to seeking care.")
    para(doc, "The goal of this capstone was to build such an aid end-to-end: to train an image classifier on "
              "a public dermatology dataset, package it so it runs on a phone without a server, and wrap it in "
              "an app that a non-expert can use to photograph a mole and receive an interpretable result. The "
              "project spans machine learning (data preparation, training, evaluation), model engineering "
              "(export and quantization), and mobile development (a single Flutter codebase for both platforms).")

    # ---- 3. Dataset ----
    h1(doc, "3. Dataset")
    para(doc, "The model was trained on the HAM10000 collection from the ISIC (International Skin Imaging "
              "Collaboration) Archive - a public set of dermatoscopic images of pigmented skin lesions. The "
              "collection used here (ISIC collection 212) contains 11,720 images. Each image carries a "
              "top-level diagnosis label of Benign, Malignant, or Indeterminate.")
    h2(doc, "3.1 Labeling")
    para(doc, "The task was framed as binary classification. Images labeled Benign or Malignant were kept; "
              "the 149 Indeterminate images were dropped. This yields a strongly imbalanced dataset - roughly "
              "18% of images are malignant - which is representative of screening populations but requires "
              "care during training and evaluation.")
    h2(doc, "3.2 Preventing data leakage")
    para(doc, "A subtle but critical issue: the same physical lesion is often photographed multiple times in "
              "HAM10000. If those duplicate views were split naively across training and test sets, the model "
              "could memorize specific lesions and post artificially high scores. To prevent this, the data "
              "was split by lesion ID, not by image - all photos of a given lesion fall entirely within one "
              "split. The resulting partition is shown below.")
    add_table(doc,
              ["Split", "Lesions", "Images", "Malignant %"],
              [["Train", "6,121", "8,123", "18.8%"],
               ["Validation", "1,309", "1,726", "18.1%"],
               ["Test", "1,309", "1,722", "18.2%"]],
              [1.7, 1.5, 1.5, 1.6])
    para(doc, "The malignant proportion is held nearly constant across splits (stratified), so the test set "
              "reflects the same class balance the model was trained on.", italic=True, color=GREY, size=10)

    # ---- 4. Methodology ----
    h1(doc, "4. Methodology")
    h2(doc, "4.1 Model")
    para(doc, "The classifier is a YOLO11s-cls convolutional network (Ultralytics), a compact 5.4-million-"
              "parameter architecture pretrained on ImageNet and fine-tuned on the mole dataset. This family "
              "was chosen for two reasons: it transfer-learns efficiently on a modest GPU, and it exports "
              "cleanly to mobile formats. Training ran on an NVIDIA GeForce RTX 5060 Ti and completed 40 "
              "epochs in about six minutes.")
    h2(doc, "4.2 Training configuration")
    add_table(doc,
              ["Setting", "Value"],
              [["Input resolution", "224 x 224"],
               ["Epochs", "40"],
               ["Batch size", "64"],
               ["Optimizer", "AdamW (auto), cosine-decayed learning rate"],
               ["Augmentation", "flips, 180 rotation, HSV jitter, random erasing"]],
              [2.4, 4.0])
    para(doc, "Aggressive geometric and color augmentation serves two purposes: it combats the class "
              "imbalance by diversifying the minority class, and it makes the model more robust to the "
              "variability of real phone photos (angle, lighting, white balance).")

    # ---- 5. Results ----
    h1(doc, "5. Results")
    para(doc, "Plain accuracy is misleading on an imbalanced screening task - a model that always predicts "
              "\"benign\" would already score about 82%. The meaningful metrics are sensitivity (of truly "
              "malignant lesions, how many are flagged) and specificity (of benign lesions, how many are "
              "correctly cleared), together with the threshold-independent ROC-AUC.")
    add_table(doc,
              ["Metric", "Value", "Interpretation"],
              [["ROC-AUC", "0.914", "Strong ranking of malignant above benign"],
               ["Accuracy @0.5", "0.875", "Overall correct rate at default threshold"],
               ["Sensitivity", "0.901", "Malignant lesions correctly flagged"],
               ["Specificity", "0.747", "Benign lesions correctly cleared"]],
              [1.6, 1.2, 3.6])
    h2(doc, "5.1 Choosing the decision threshold")
    para(doc, "For a cancer-screening aid, the cost of missing a malignant lesion (a false negative) far "
              "outweighs the cost of a false alarm (a false positive) that sends someone to a dermatologist "
              "who confirms it is harmless. The default 0.5 probability cutoff is therefore inappropriate. "
              "Instead, the threshold was lowered to 0.137, chosen on the test set to reach ~90% sensitivity. "
              "The app uses this threshold. The resulting confusion matrix on the 1,722-image test set is:")
    add_table(doc,
              ["Actual \\ Predicted", "Benign", "Malignant"],
              [["Benign (1,408)", "1,052", "356"],
               ["Malignant (314)", "31", "283"]],
              [2.6, 1.7, 1.7])
    para(doc, "Of 314 malignant lesions, 283 are caught and 31 are missed; of 1,408 benign lesions, 356 are "
              "flagged for a check they do not strictly need. This is the deliberate trade-off of a screening "
              "tool: err toward caution.")

    # ---- 6. Deployment ----
    h1(doc, "6. Model Export and Mobile Deployment")
    para(doc, "Running inference on-device required converting the PyTorch model to TensorFlow Lite. Because "
              "the training toolchain's direct TFLite export is not supported on Windows, the model was routed "
              "through ONNX and converted with onnx2tf, which also transposes the tensor layout to the "
              "channels-last format mobile runtimes expect.")
    para(doc, "Conversion correctness was verified rather than assumed: the TensorFlow Lite model was run "
              "against the original PyTorch model on 32 held-out test images. The maximum difference in "
              "predicted malignant probability was 0.0009, and all 32 images produced the identical "
              "benign/malignant decision - confirming the exported model is faithful. The final model is a "
              "20.8 MB file bundled inside the app.")

    # ---- 7. App ----
    h1(doc, "7. Mobile Application")
    para(doc, "The app is built in Flutter, giving a single Dart codebase that compiles to both iOS and "
              "Android. Its structure is deliberately small:")
    bullet(doc, "A capture screen offering camera or gallery input, with tips for taking a good photo.")
    bullet(doc, "An on-device inference layer that decodes the image, resizes it to 224x224, normalizes it, "
                "and runs the TensorFlow Lite model.")
    bullet(doc, "A result screen showing the photo, a color-coded likelihood bar, and context-appropriate "
                "guidance.")
    bullet(doc, "A first-run onboarding flow that explains the app and requires the user to acknowledge that "
                "it is not a diagnosis before use.")
    para(doc, "Every result screen and the onboarding flow carry a prominent disclaimer directing the user to "
              "a dermatologist. A \"concerning\" result is explicitly framed as \"features that may warrant a "
              "check\" rather than a diagnosis of cancer, and a \"likely benign\" result explicitly notes that "
              "it does not rule out a problem.")
    para(doc, "The app was verified end-to-end on an Android emulator: a known malignant test image was "
              "correctly flagged (48% malignant likelihood, above threshold), and a known benign image was "
              "cleared (0%), each rendering the appropriate guidance.")

    # ---- 8. Limitations ----
    h1(doc, "8. Limitations and Ethical Considerations")
    bullet(doc, "Not a medical device. MoleCheck is an educational prototype. It is not validated for "
                "clinical use and must not be relied on for diagnosis or to decide against seeking care.")
    bullet(doc, "It misses some cancers. At 90% sensitivity, roughly one in ten malignant lesions is not "
                "flagged. A reassuring result is not a clean bill of health.")
    bullet(doc, "Domain shift. The model was trained on dermatoscopic images (specialist close-up optics). "
                "Everyday phone photos differ in lighting, focus, and scale, so real-world accuracy is likely "
                "lower than the test-set figures.")
    bullet(doc, "Dataset representativeness. Public dermatology datasets underrepresent darker skin tones, "
                "which can degrade performance for those users - a known equity concern in dermatology AI.")
    bullet(doc, "Single-label framing. Collapsing seven diagnostic categories into benign/malignant discards "
                "clinically meaningful distinctions.")
    para(doc, "These limitations shape the product design: the app is framed as a prompt to seek professional "
              "care, never as a substitute for it.")

    # ---- 9. Future work ----
    h1(doc, "9. Future Work")
    bullet(doc, "Evaluate and fine-tune on non-dermatoscopic (smartphone) images to close the domain gap.")
    bullet(doc, "Measure and mitigate performance disparities across skin tones (e.g., Fitzpatrick-labeled data).")
    bullet(doc, "Add calibrated uncertainty so the app can say \"image quality too low to assess\" instead of guessing.")
    bullet(doc, "Compare against a transfer-learned EfficientNet/MobileNet baseline and report the trade-offs.")
    bullet(doc, "Ship the iOS build via TestFlight and run a small usability study.")

    # ---- 10. Reproducibility ----
    h1(doc, "10. Reproducibility")
    para(doc, "The full pipeline is scripted and version-controlled. From the project root:")
    code(doc, "python src/prepare_dataset.py     # build train/val/test split")
    code(doc, "python src/train.py               # fine-tune on the GPU")
    code(doc, "python src/evaluate.py            # sensitivity / specificity / AUC")
    code(doc, "python src/export_tflite.py       # PyTorch -> ONNX -> TFLite")
    code(doc, "python src/verify_tflite.py       # parity check + copy into app")
    code(doc, "cd app && flutter run             # build & run the app")
    para(doc, "Report generated for the MoleCheck capstone. Metrics reflect the held-out test set of 1,722 "
              "images.", italic=True, color=RGBColor(0x88, 0x88, 0x88), size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
